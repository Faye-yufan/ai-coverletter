import openai
import re
import os
import json
from fpdf import FPDF
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx2pdf import convert
import PyPDF2

# Set up OpenAI API credentials
OPENAI_API_KEY = "your-api-key"
openai.api_key = OPENAI_API_KEY

def read_docx(file_path):
    # doc = docx.Document(file_path)
    full_text = []
    # for para in doc.paragraphs:
    #     full_text.append(para.text)
    # Open the PDF file in read-binary mode
    with open(file_path, 'rb') as file:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(file)

        # Get the total number of pages in the PDF file
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            full_text.append(page.extract_text())
    return '\n'.join(full_text)

def generate_response(prompts):
    ans = ""
    response = openai.ChatCompletion.create(
        model = model_id,
        messages = prompts
    )
    ans = response['choices'][0].message.content

    return ans

def create_cover_letter_doc(cover_letter, company_info, output_file):
    doc = Document()

    def add_paragraph_with_font(text):
        paragraph = doc.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)  # Set the font size, if needed
        return paragraph

    # Add your name and address (modify this as needed)
    name_address = 'Your name and address'
    add_paragraph_with_font(name_address)

    # Add the recipient's address
    recipient_address = company_info
    add_paragraph_with_font(recipient_address)

    # Add the date
    from datetime import date
    today = date.today().strftime("%B %d, %Y")
    add_paragraph_with_font(today)

    # Add the cover letter text
    add_paragraph_with_font(cover_letter)

    # Save the Word document
    doc.save(output_file)

def user_input_conversion(file_name):
    user_input = input("Enter '1' to convert a .docx file to .pdf: ")

    if user_input == '1':
        input_docx = f'{file_name}.docx'  
        output_pdf = f'{file_name}.pdf'  

        # Ensure input file is a .docx file
        if not input_docx.lower().endswith('.docx'):
            print(f"{input_docx} is not a .docx file.")
            return

        # Check if input file exists
        if not os.path.isfile(input_docx):
            print(f"{input_docx} does not exist.")
            return

        # Convert the .docx file to a .pdf file
        convert(input_docx, output_pdf)

        print(f"{input_docx} has been converted to {output_pdf}.")
    return

# Extract relevant information from the job description and your resume
company_name = "LOCKHEED MARTIN CORPORATION"
jd_file_name = 'jd.txt'

with open(jd_file_name, 'r') as file:
    job_description = file.read()

resume = read_docx("your-resume.pdf")

model_id = 'gpt-3.5-turbo'

skills_section = """
Languages/Technologies: ...
Others: ...
"""
is_finish_resume = 0
while is_finish_resume != "1":
    prompts = [
        {
            "role": "system",
            "content": f"Please rewrite the following skills section of my resume based on the job description provided.\n\n \
                    Job description:\n{job_description}\n\nCurrent skills section:\n{skills_section}\n\nRewritten skills section \
                        please only return 2 lines, one lines for Languages/Technologies and the second line for Others ",
        }
    ]
    new_skills = generate_response(prompts)
    print("---------------------------------")
    print(new_skills)
    print("---------------------------------")
    is_finish_resume = input("Enter '1' if you finished resume editing and want to generate coverletter: ")

prompts = [{'role': 'system', 'content': f"Write a cover letter for the following job description: '{job_description}' based on my resume: '{resume}'"},
           {'role': 'user', 'content': "Rewrite it more conversational."},
           {'role': 'user', 'content': "Rewrite it again but use shorter paragraphs, and less than 270 words."}]
cover_letter = generate_response(prompts)

prompts = [{'role': 'system', 'content': f"Can you give me the company name, Company Address, City, State, Zip Code for '{company_name}'"},
           {'role': 'user', 'content': "Please give just 3 lines, one line only contains company name, one line only contains Company address, one line City, State, Zip Code"}]
company_info = generate_response(prompts)

output_file_name = f'Cover_letter_{company_name}'
create_cover_letter_doc(cover_letter, company_info, f'{output_file_name}.docx')

user_input_conversion(output_file_name)
